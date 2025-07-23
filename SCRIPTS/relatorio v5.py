import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# Caminho do arquivo Excel
arquivo_excel = "itens.xlsx"

# Verifica se o arquivo existe
if not os.path.exists(arquivo_excel):
    raise FileNotFoundError(f"Arquivo '{arquivo_excel}' não encontrado na pasta do script.")

# Lê os dados
try:
    df_gastos = pd.read_excel(arquivo_excel, sheet_name="Gastos")
    df_vendas = pd.read_excel(arquivo_excel, sheet_name="Vendas")
except Exception as e:
    raise ValueError(f"Erro ao ler planilhas: {e}")

# Processamento de datas
df_gastos["data"] = pd.to_datetime(df_gastos["data"])
df_gastos["semana"] = df_gastos["data"] - pd.to_timedelta((df_gastos["data"].dt.weekday + 1) % 7, unit="D")
df_gastos["semana"] = df_gastos["semana"].dt.date
df_gastos["mes"] = df_gastos["data"].dt.to_period("M")

df_vendas["data"] = pd.to_datetime(df_vendas["data"])
df_vendas["semana"] = df_vendas["data"] - pd.to_timedelta((df_vendas["data"].dt.weekday + 1) % 7, unit="D")
df_vendas["semana"] = df_vendas["semana"].dt.date
df_vendas["mes"] = df_vendas["data"].dt.to_period("M")

# Agrupamentos
gastos_semanais = df_gastos.groupby("semana")["valor"].sum()
vendas_semanais = df_vendas.groupby("semana")["valor"].sum()
lucro_semanais = vendas_semanais.subtract(gastos_semanais, fill_value=0)

gastos_mensais = df_gastos.groupby("mes")["valor"].sum()
vendas_mensais = df_vendas.groupby("mes")["valor"].sum()
lucro_mensais = vendas_mensais.subtract(gastos_mensais, fill_value=0)

# Gera gráfico
plt.figure(figsize=(10, 6))
x = range(len(gastos_semanais))
plt.bar(x, gastos_semanais.values, width=0.3, label="Gastos", align="center")
plt.bar([i + 0.3 for i in x], vendas_semanais.values, width=0.3, label="Vendas", align="center")
plt.xticks([i + 0.15 for i in x], gastos_semanais.index, rotation=45)
plt.title("Gastos vs Vendas por Semana")
plt.legend()
plt.tight_layout()

grafico_path = os.path.abspath("grafico_semanal.png")
plt.savefig(grafico_path)
plt.close()

# Cria o relatório
doc = Document()
doc.add_heading("Relatório Financeiro", 0)
doc.add_picture(grafico_path, width=Inches(6))

doc.add_heading("Resumo Semanal", level=1)
tabela = doc.add_table(rows=1, cols=4)
hdr = tabela.rows[0].cells
hdr[0].text = "Semana"
hdr[1].text = "Vendas (R$)"
hdr[2].text = "Gastos (R$)"
hdr[3].text = "Lucro (R$)"

for semana in sorted(set(gastos_semanais.index).union(vendas_semanais.index)):
    row = tabela.add_row().cells
    row[0].text = str(semana)
    row[1].text = f"{vendas_semanais.get(semana, 0):,.2f}"
    row[2].text = f"{gastos_semanais.get(semana, 0):,.2f}"
    row[3].text = f"{lucro_semanais.get(semana, 0):,.2f}"

doc.add_heading("Resumo Mensal", level=1)
tabela = doc.add_table(rows=1, cols=4)
hdr = tabela.rows[0].cells
hdr[0].text = "Mês"
hdr[1].text = "Vendas (R$)"
hdr[2].text = "Gastos (R$)"
hdr[3].text = "Lucro (R$)"

for mes in lucro_mensais.index:
    row = tabela.add_row().cells
    row[0].text = str(mes)
    row[1].text = f"{vendas_mensais.get(mes, 0):,.2f}"
    row[2].text = f"{gastos_mensais.get(mes, 0):,.2f}"
    row[3].text = f"{lucro_mensais.get(mes, 0):,.2f}"

# Salva o documento
relatorio_path = os.path.abspath("relatorio_completo.docx")
doc.save(relatorio_path)

print(f"✅ Relatório salvo com sucesso em: {relatorio_path}")
